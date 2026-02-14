"""
Document Export - Export chat data to various formats.
Supports: TXT, JSON, PDF, XML, DOCX, ODT, HTML
Also supports summary generation and printing.
"""
import io
import os
import json
import time
import logging
from pathlib import Path
from typing import Dict, Any, List
logger = logging.getLogger("DocumentExporter")

class DocumentExporter:
    """Exports chat data and summaries to multiple document formats."""
    def __init__(self, config: dict):
        self.config = config
        self.exports_path = Path(config.get('exports_path', 'data/exports'))
        self.exports_path.mkdir(parents=True, exist_ok=True)
    def export(self, format: str, chat_data: Dict[str, Any]) -> Dict[str, Any]:
        """Export chat to the specified format."""
        exporters = {
            'txt': self._export_txt,
            'json': self._export_json,
            'pdf': self._export_pdf,
            'xml': self._export_xml,
            'docx': self._export_docx,
            'odt': self._export_odt,
            'html': self._export_html,
        }
        exporter = exporters.get(format)
        if not exporter:
            return {"status": "error", "message": f"Unsupported export format: {format}"}
        try:
            character = chat_data.get('character', 'Unknown')
            timestamp = time.strftime('%Y%m%d_%H%M%S')
            base_filename = f"chat_{character}_{timestamp}"
            filepath = exporter(chat_data, base_filename)
            filename = os.path.basename(filepath)
            return {
                "status": "ok",
                "filename": filename,
                "download_url": f"/documents/download/{filename}",
                "format": format
            }
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return {"status": "error", "message": str(e)}
    def _format_messages(self, chat_data: Dict) -> List[Dict]:
        """Normalize message format."""
        messages = chat_data.get('messages', [])
        return [
            {
                "name": m.get('name', m.get('role', 'Unknown')),
                "role": m.get('role', 'unknown'),
                "content": m.get('content', ''),
            }
            for m in messages if m.get('content', '').strip()
        ]
    def _export_txt(self, chat_data: Dict, base_filename: str) -> str:
        messages = self._format_messages(chat_data)
        filepath = self.exports_path / f"{base_filename}.txt"
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"Chat Export - {chat_data.get('character', 'Unknown')}\n")
            f.write(f"Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 60 + "\n\n")
            for msg in messages:
                f.write(f"[{msg['name']}]:\n")
                f.write(f"{msg['content']}\n\n")
                f.write("-" * 40 + "\n\n")
        return str(filepath)
    def _export_json(self, chat_data: Dict, base_filename: str) -> str:
        filepath = self.exports_path / f"{base_filename}.json"
        export_data = {
            "character": chat_data.get('character', 'Unknown'),
            "export_date": time.strftime('%Y-%m-%d %H:%M:%S'),
            "messages": self._format_messages(chat_data)
        }
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        return str(filepath)
    def _export_pdf(self, chat_data: Dict, base_filename: str) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        filepath = self.exports_path / f"{base_filename}.pdf"
        doc = SimpleDocTemplate(str(filepath), pagesize=A4,
                                topMargin=50, bottomMargin=50)
        styles = getSampleStyleSheet()
        user_style = ParagraphStyle(
            'UserMsg', parent=styles['Normal'],
            backColor=HexColor('#e3f2fd'),
            spaceBefore=6, spaceAfter=6,
            leftIndent=10, rightIndent=10
        )
        ai_style = ParagraphStyle(
            'AIMsg', parent=styles['Normal'],
            backColor=HexColor('#f5f5f5'),
            spaceBefore=6, spaceAfter=6,
            leftIndent=10, rightIndent=10
        )
        name_style = ParagraphStyle(
            'NameStyle', parent=styles['Normal'],
            fontName='Helvetica-Bold', fontSize=10
        )
        story = []
        story.append(Paragraph(
            f"Chat Export - {chat_data.get('character', 'Unknown')}", styles['Title']
        ))
        story.append(Paragraph(
            f"Exported: {time.strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']
        ))
        story.append(Spacer(1, 20))
        messages = self._format_messages(chat_data)
        for msg in messages:
            style = user_style if msg['role'] == 'user' else ai_style
            # Escape XML special chars for reportlab
            content = (msg['content']
                       .replace('&', '&amp;')
                       .replace('<', '&lt;')
                       .replace('>', '&gt;'))
            name = (msg['name']
                    .replace('&', '&amp;')
                    .replace('<', '&lt;')
                    .replace('>', '&gt;'))
            story.append(Paragraph(f"<b>{name}</b>:", name_style))
            story.append(Paragraph(content, style))
            story.append(Spacer(1, 8))
        doc.build(story)
        return str(filepath)
    def _export_xml(self, chat_data: Dict, base_filename: str) -> str:
        from lxml import etree
        root = etree.Element("chat_export")
        root.set("character", chat_data.get('character', 'Unknown'))
        root.set("date", time.strftime('%Y-%m-%d %H:%M:%S'))
        messages_el = etree.SubElement(root, "messages")
        for msg in self._format_messages(chat_data):
            msg_el = etree.SubElement(messages_el, "message")
            msg_el.set("role", msg['role'])
            name_el = etree.SubElement(msg_el, "name")
            name_el.text = msg['name']
            content_el = etree.SubElement(msg_el, "content")
            content_el.text = msg['content']
        filepath = self.exports_path / f"{base_filename}.xml"
        tree = etree.ElementTree(root)
        tree.write(str(filepath), pretty_print=True, xml_declaration=True,
                   encoding='utf-8')
        return str(filepath)
    def _export_docx(self, chat_data: Dict, base_filename: str) -> str:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        # Title
        title = doc.add_heading(
            f"Chat Export - {chat_data.get('character', 'Unknown')}", level=1
        )
        doc.add_paragraph(f"Exported: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")  # spacer
        for msg in self._format_messages(chat_data):
            # Name
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(f"{msg['name']}:")
            name_run.bold = True
            name_run.font.size = Pt(11)
            if msg['role'] == 'user':
                name_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            else:
                name_run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
            # Content
            content_para = doc.add_paragraph(msg['content'])
            content_para.paragraph_format.space_after = Pt(12)
        filepath = self.exports_path / f"{base_filename}.docx"
        doc.save(str(filepath))
        return str(filepath)
    def _export_odt(self, chat_data: Dict, base_filename: str) -> str:
        from odf.opendocument import OpenDocumentText
        from odf.text import P, H
        from odf import style as odf_style
        doc = OpenDocumentText()
        # Create styles
        bold_style = odf_style.Style(name="Bold", family="text")
        bold_prop = odf_style.TextProperties(fontweight="bold")
        bold_style.addElement(bold_prop)
        doc.automaticstyles.addElement(bold_style)
        # Title
        h = H(outlinelevel=1, text=f"Chat Export - {chat_data.get('character', 'Unknown')}")
        doc.text.addElement(h)
        p = P(text=f"Exported: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.text.addElement(p)
        doc.text.addElement(P(text=""))
        for msg in self._format_messages(chat_data):
            name_p = P(text=f"{msg['name']}:")
            doc.text.addElement(name_p)
            content_p = P(text=msg['content'])
            doc.text.addElement(content_p)
            doc.text.addElement(P(text=""))
        filepath = self.exports_path / f"{base_filename}.odt"
        doc.save(str(filepath))
        return str(filepath)
    def _export_html(self, chat_data: Dict, base_filename: str) -> str:
        messages = self._format_messages(chat_data)
        def esc(text):
            return (text.replace('&', '&amp;')
                    .replace('<', '&lt;')
                    .replace('>', '&gt;')
                    .replace('\n', '<br>'))
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Chat Export - {esc(chat_data.get('character', 'Unknown'))}</title>
<style>
    body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 800px;
           margin: 0 auto; padding: 20px; background: #fafafa; }}
    h1 {{ color: #333; }}
    .message {{ margin: 12px 0; padding: 12px 16px; border-radius: 12px; }}
    .user {{ background: #e3f2fd; border-left: 4px solid #1a73e8; }}
    .assistant {{ background: #f5f5f5; border-left: 4px solid #666; }}
    .system {{ background: #fff3e0; border-left: 4px solid #ff9800; }}
    .name {{ font-weight: bold; margin-bottom: 4px; font-size: 0.9em; }}
    .meta {{ color: #999; font-size: 0.8em; margin-bottom: 10px; }}
</style>
</head>
<body>
<h1>Chat Export - {esc(chat_data.get('character', 'Unknown'))}</h1>
<p class="meta">Exported: {time.strftime('%Y-%m-%d %H:%M:%S')}</p>
<hr>
"""
        for msg in messages:
            role_class = msg['role']
            html += f"""<div class="message {role_class}">
    <div class="name">{esc(msg['name'])}</div>
    <div class="content">{esc(msg['content'])}</div>
</div>
"""
        html += "</body></html>"
        filepath = self.exports_path / f"{base_filename}.html"
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html)
        return str(filepath)
    def generate_summary(self, chat_data: Dict[str, Any]) -> str:
        """Generate a text summary of the chat.
        Uses extractive summarization (no LLM dependency).
        """
        messages = self._format_messages(chat_data)
        if not messages:
            return "No messages to summarize."
        character = chat_data.get('character', 'Unknown')
        total_msgs = len(messages)
        user_msgs = sum(1 for m in messages if m['role'] == 'user')
        ai_msgs = total_msgs - user_msgs
        # Extract key topics by frequency
        from collections import Counter
        import re
        all_text = ' '.join(m['content'] for m in messages)
        words = re.findall(r'\b[a-zA-Z]{4,}\b', all_text.lower())
        # Filter common stop words
        stop_words = {
            'that', 'this', 'with', 'have', 'from', 'they', 'been',
            'would', 'could', 'should', 'their', 'there', 'about',
            'which', 'when', 'what', 'will', 'your', 'just', 'like',
            'know', 'think', 'well', 'also', 'then', 'than', 'more',
            'some', 'very', 'into', 'over', 'after', 'only', 'other',
            'were', 'them', 'being', 'does', 'much', 'said', 'each'
        }
        filtered = [w for w in words if w not in stop_words]
        common = Counter(filtered).most_common(10)
        topics = ', '.join(word for word, _ in common)
        # First and last exchanges
        first_user = next((m for m in messages if m['role'] == 'user'), None)
        last_user = next((m for m in reversed(messages) if m['role'] == 'user'), None)
        summary_parts = [
            f"Conversation Summary with {character}",
            f"Total messages: {total_msgs} ({user_msgs} user, {ai_msgs} AI)",
            f"Key topics: {topics}",
        ]
        if first_user:
            opening = first_user['content'][:150]
            summary_parts.append(f"Conversation opened with: \"{opening}...\"")
        if last_user and last_user != first_user:
            closing = last_user['content'][:150]
            summary_parts.append(f"Last user message: \"{closing}...\"")
        # Content length stats
        total_chars = sum(len(m['content']) for m in messages)
        avg_len = total_chars // max(total_msgs, 1)
        summary_parts.append(
            f"Total content length: ~{total_chars // 1000}K characters "
            f"(avg {avg_len} chars/message)"
        )
        return '\n'.join(summary_parts)
